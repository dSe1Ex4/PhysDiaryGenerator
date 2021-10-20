"""
@author Аномальная Субстанция

Инструкция:
 1.Указать даты (dates) и параметры в Настройках
 2.Запустить скрипт
 3.Следуй иструкциям
 4.Файл будет в папке со скриптом <ваше имя>.docx
"""


from docx import Document
import random

"""Настройка"""
dates = [
    "15.09.2021",
    "22.09.2021",
    "29.09.2021",
]

# Указать ваш НОРМАЛЬНЫЙ пульс
pulseMin = TODO(60)  # Ваш минимальный пульс
pulseMax = TODO(70)  # Ваш максимальный пульс

# Указать ваш НОРМАЛЬНЫЙ ЧД
chDeMin = TODO(40)  # Ваша минимальная ЧД
chDeMax = TODO(50)  # Ваша максимальная ЧД

# Указать на сколько у вас меняется пульс (ВАШ_НОРМАЛЬНЫЙ_ПАРАМЕТР - ПАРАМЕТР_ПОСЛЕ_ТРЕНИ)
diffPulseMin = TODO(30)
diffPulseMax = TODO(50)

# Указать на сколько у вас меняется ЧД (ВАШ_НОРМАЛЬНЫЙ_ПАРАМЕТР - ПАРАМЕТР_ПОСЛЕ_ТРЕНИ)
diffChDeMin = TODO(30)
diffChDeMax = TODO(40)
"""Настройка закончена"""

iBPM = random.randrange(pulseMin, pulseMax)
iB = random.randrange(chDeMin, chDeMax)


doc=Document()

name = input("Ваше ФИО>> ")
doc.add_heading(name, 0)

table = doc.add_table(rows=len(dates)+1, cols=5)
table.style = 'Table Grid'

row = table.rows[0]
row.cells[0].add_paragraph("Дата занятия", style="Heading 5")
row.cells[1].add_paragraph("ЧСС до", style="Heading 5")
row.cells[2].add_paragraph("ЧД до", style="Heading 5")
row.cells[3].add_paragraph("ЧСС после", style="Heading 5")
row.cells[4].add_paragraph("ЧД после", style="Heading 5")

tbl_index = 0
tbl_rowindex = 1

ind = 0
inp = dates[ind]


while inp != "":
    iBPM = random.randrange(pulseMin, pulseMax)
    iB = random.randrange(chDeMin, chDeMax)

    row = table.rows[tbl_rowindex]

    row.cells[0].add_paragraph(str(inp))
    tbl_index += 1
    row.cells[1].add_paragraph(str(iBPM))
    tbl_index += 1
    row.cells[2].add_paragraph(str(iB))
    tbl_index += 1
    row.cells[3].add_paragraph(str(iBPM + random.randrange(diffPulseMin, diffPulseMax)))
    tbl_index += 1
    row.cells[4].add_paragraph(str(iB + random.randrange(diffChDeMin, diffChDeMax)))
    tbl_index += 1
    tbl_rowindex += 1
    # print(str(ind)+" Added: "+inp)
    # print(inp+": "+str(iBPM)+" "+str(iB)+" "+str(iBPM + random.randrange(29, 54))+" "+str(iB + random.randrange(29, 42)))
    ind += 1
    if ind <= len(dates)-1:
        inp = dates[ind]
    else:
        inp = ""

doc.save(name+'.docx')
